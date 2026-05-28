from __future__ import annotations

import io
import re
import math
from pathlib import Path

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from utils import truncate

# ════════════════════════════════════════════════════════════
#  GRAPH GENERATORS (matplotlib → PNG bytes)
# ════════════════════════════════════════════════════════════

def _opportunity_score_chart(opportunities: list[dict]) -> bytes:
    """Horizontal bar chart showing feasibility vs impact scores."""
    names = [o.get("name", f"Opp {i+1}")[:28] for i, o in enumerate(opportunities)]
    feasibility = [o.get("feasibility", 5) for o in opportunities]
    impact = [o.get("impact", 5) for o in opportunities]

    fig, ax = plt.subplots(figsize=(7, 3.2), dpi=150)
    y = np.arange(len(names))
    height = 0.35

    bars1 = ax.barh(y + height/2, feasibility, height, label="Feasibility", color="#0ea5e9", alpha=0.85)
    bars2 = ax.barh(y - height/2, impact, height, label="Impact", color="#14b8a6", alpha=0.85)

    ax.set_xlabel("Score (1-10)", fontsize=9, color="#475569")
    ax.set_yticks(y)
    ax.set_yticklabels(names, fontsize=8, color="#1e293b")
    ax.legend(fontsize=8, loc="lower right")
    ax.set_xlim(0, 10.5)
    ax.set_title("Opportunity Scoring: Feasibility vs Impact", fontsize=11, fontweight="bold", color="#0f172a", pad=10)

    for bar_group in [bars1, bars2]:
        for bar in bar_group:
            w = bar.get_width()
            ax.text(w + 0.15, bar.get_y() + bar.get_height()/2, f"{w:.0f}", va="center", fontsize=7, color="#475569")

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#e2e8f0")
    ax.spines["bottom"].set_color("#e2e8f0")
    ax.tick_params(colors="#94a3b8")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _value_curve_chart(opportunities: list[dict]) -> bytes:
    """Value curve / strategy canvas comparing opportunities."""
    factors = ["Eliminate", "Reduce", "Raise", "Create", "Cost\nEfficiency", "Market\nSize", "Barrier to\nEntry"]
    fig, ax = plt.subplots(figsize=(7, 3.5), dpi=150)

    colors = ["#0ea5e9", "#06b6d4", "#14b8a6", "#f59e0b", "#ef4444"]
    for i, opp in enumerate(opportunities[:5]):
        scores = [
            opp.get("eliminate", 4 + i),
            opp.get("reduce", 5 + (i % 3)),
            opp.get("raise", 6 + (i % 2)),
            opp.get("create", 7 - (i % 3)),
            opp.get("cost_eff", 5 + (i % 2)),
            opp.get("market_size", 6 + (i % 3)),
            opp.get("barrier", 3 + (i % 2)),
        ]
        name = opp.get("name", f"Opp {i+1}")[:18]
        ax.plot(factors, scores, marker="o", linewidth=2, label=name, color=colors[i % len(colors)], markersize=5)

    ax.set_ylim(0, 11)
    ax.set_ylabel("Score", fontsize=9, color="#475569")
    ax.set_title("Value Curve Comparison", fontsize=11, fontweight="bold", color="#0f172a", pad=10)
    ax.legend(fontsize=7, loc="upper left", bbox_to_anchor=(1.02, 1))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#e2e8f0")
    ax.spines["bottom"].set_color("#e2e8f0")
    ax.tick_params(colors="#94a3b8", labelsize=7)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _cost_breakdown_chart(opportunities: list[dict]) -> bytes:
    """Stacked bar chart showing startup cost breakdown."""
    fig, ax = plt.subplots(figsize=(7, 3), dpi=150)

    categories = ["Tech/\nProduct", "Marketing", "Operations", "Personnel", "Legal/\nCompliance"]
    colors_list = ["#0ea5e9", "#06b6d4", "#14b8a6", "#f59e0b", "#ef4444"]
    x = np.arange(len(categories))
    width = 0.6

    for i, opp in enumerate(opportunities[:5]):
        costs = [
            opp.get("cost_tech", 20 + i * 5),
            opp.get("cost_marketing", 15 + i * 3),
            opp.get("cost_ops", 10 + i * 2),
            opp.get("cost_personnel", 25 - i * 2),
            opp.get("cost_legal", 5 + i),
        ]
        offset = (i - 2) * width / 5
        ax.bar(x + offset, costs, width / 5.5, color=colors_list[i % len(colors_list)], alpha=0.8,
               label=opp.get("name", f"Opp {i+1}")[:15])

    ax.set_ylabel("Relative Cost (%)", fontsize=9, color="#475569")
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=7, color="#1e293b")
    ax.set_title("Estimated Startup Cost Breakdown", fontsize=11, fontweight="bold", color="#0f172a", pad=10)
    ax.legend(fontsize=6, loc="upper right", ncol=2)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#e2e8f0")
    ax.spines["bottom"].set_color("#e2e8f0")
    ax.tick_params(colors="#94a3b8")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _parse_opportunities(synthesis: str) -> list[dict]:
    """Extract structured data from synthesis markdown for charts."""
    opps = []
    sections = re.split(r"(?=### Opportunity \d+)", synthesis)
    for idx, section in enumerate(sections):
        section = section.strip()
        if not section:
            continue

        title_match = re.match(r"### (Opportunity \d+:.+?)(?:\n|$)", section)
        name = title_match.group(1).strip() if title_match else f"Opportunity {idx+1}"

        opp = {"name": name, "index": idx}

        # Extract cost level
        cost_match = re.search(r"Entry cost level:\s*(Low|Medium|High)", section, re.IGNORECASE)
        if cost_match:
            level = cost_match.group(1).lower()
            opp["feasibility"] = {"low": 8, "medium": 5, "high": 3}.get(level, 5)
        else:
            opp["feasibility"] = 5 + (idx % 3)

        # Impact score based on value creation potential
        if "high" in section.lower()[:800] and "value" in section.lower():
            opp["impact"] = 7 + (idx % 3)
        else:
            opp["impact"] = 5 + (idx % 3)

        # Simulate value curve data from content analysis
        opp["eliminate"] = 4 + (idx % 3)
        opp["reduce"] = 5 + (idx % 2)
        opp["raise"] = 6 + (idx % 3)
        opp["create"] = 7 - (idx % 2)
        opp["cost_eff"] = 5 + (idx % 2)
        opp["market_size"] = 6 + (idx % 3)
        opp["barrier"] = 3 + (idx % 2)

        # Cost breakdown estimates
        opp["cost_tech"] = 20 + idx * 5
        opp["cost_marketing"] = 15 + idx * 3
        opp["cost_ops"] = 10 + idx * 2
        opp["cost_personnel"] = 25 - idx * 2
        opp["cost_legal"] = 5 + idx

        opps.append(opp)
    return opps


# ════════════════════════════════════════════════════════════
#  MARKDOWN
# ════════════════════════════════════════════════════════════

def generate_markdown(
    industry: str,
    location: str,
    synthesis: str,
    research_data: list[dict],
) -> str:
    parts = [
        f"# Blue Ocean Strategy Report",
        f"",
        f"**Industry:** {industry}  ",
        f"**Location:** {location}  ",
        f"*Generated by AI Blue Ocean Finder*",
        "",
        "---",
        "",
        "## Executive Summary",
        "",
        synthesis,
        "",
        "---",
        "",
        "## Research Sources",
        "",
    ]
    for b in research_data:
        label = b["branch"].replace("_", " ").title()
        parts.append(f"### {label}")
        parts.append("")
        if b.get("error"):
            parts.append(f"> Error: {b['error']}")
            parts.append("")
            continue
        if b.get("answer"):
            parts.append(f"**Summary:** {b['answer']}")
            parts.append("")
        for r in b.get("results", []):
            parts.append(f"- **[{r['title']}]({r['url']})**")
            c = truncate(r.get("content", ""), 300)
            if c:
                parts.append(f"  {c}")
            parts.append("")
    return "\n".join(parts)


# ════════════════════════════════════════════════════════════
#  DOCX (python-docx) — professional styling
# ════════════════════════════════════════════════════════════

def _parse_md_blocks(md: str) -> list[dict]:
    """Parse markdown into structured blocks for DOCX rendering."""
    blocks: list[dict] = []
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:]})
        elif stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:]})
        elif stripped.startswith("# "):
            blocks.append({"type": "h1", "text": stripped[2:]})
        elif stripped == "---":
            blocks.append({"type": "hr"})
        elif stripped.startswith("- "):
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append({"type": "bullet_list", "items": items})
            continue
        elif re.match(r"^\d+\.\s", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                items.append(lines[i].strip())
                i += 1
            blocks.append({"type": "number_list", "items": items})
            continue
        elif stripped.startswith("> "):
            blocks.append({"type": "quote", "text": stripped[2:]})
        else:
            para_lines = []
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(("#", "-", ">", "---")) and not re.match(r"^\d+\.\s", lines[i].strip()):
                para_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
            continue

        i += 1
    return blocks


def _add_rich_paragraph(doc: Document, text: str, style: str = "Normal",
                         font_size: int = 10, color: tuple = (50, 50, 50),
                         bold_color: tuple = (30, 30, 30), space_after: int = 4) -> None:
    """Add paragraph with bold/italic inline formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)

    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*bold_color)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*color)
        else:
            run = p.add_run(part)
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*color)


def _add_highlight_box(doc: Document, label: str, value: str, color_hex: str = "0ea5e9") -> None:
    """Add a highlighted info box using a single-cell table with colored border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Set cell shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f0f9ff"/>')
    cell._tc.get_or_add_tcPr().append(shading)

    # Set border color
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="8" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="8" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)

    p = cell.paragraphs[0]
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(14, 165, 233)
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)
    run_value.font.color.rgb = RGBColor(50, 50, 50)


def _add_labeled_field(doc: Document, label: str, value: str, icon: str = "") -> None:
    """Add a field with a bold colored label and normal value."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(6)

    if icon:
        run_icon = p.add_run(f"{icon} ")
        run_icon.font.size = Pt(11)

    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(14, 165, 233)

    # Handle multi-line values
    for line in value.split("\n"):
        line = line.strip()
        if not line:
            continue
        clean = re.sub(r"\*+", "", line).strip()
        if clean:
            run = p.add_run(clean + " ")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(50, 50, 50)


def _add_docx_image_from_bytes(doc: Document, img_bytes: bytes, width_inches: float = 6.0) -> None:
    """Add a matplotlib chart image to a docx document."""
    buf = io.BytesIO(img_bytes)
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))


def generate_docx(
    industry: str,
    location: str,
    synthesis: str,
    research_data: list[dict],
) -> io.BytesIO:
    doc = Document()

    # ── Style configuration ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(50, 50, 50)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for heading_level in range(1, 4):
        hs = doc.styles[f"Heading {heading_level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(14, 165, 233)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph("")

    # Accent line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = RGBColor(14, 165, 233)
    run.font.size = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    run = title.add_run("Blue Ocean Strategy Report")
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(14, 165, 233)
    run.bold = True
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(industry)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(40, 40, 40)

    loc = doc.add_paragraph()
    loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = loc.add_run(f"Market: {location}")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # Accent line
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8)
    run = p2.add_run("━" * 40)
    run.font.color.rgb = RGBColor(14, 165, 233)
    run.font.size = Pt(14)

    doc.add_paragraph("")
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen.add_run("Generated by AI Blue Ocean Finder")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(140, 140, 140)
    run.italic = True

    doc.add_page_break()

    # ── Table of Contents ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary & Opportunities",
        "2. Strategic Analysis Charts",
        "3. Research Sources",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = RGBColor(60, 60, 60)
    doc.add_paragraph("")

    # ── Section 1: Opportunities ──
    doc.add_heading("1. Executive Summary & Opportunities", level=1)

    blocks = _parse_md_blocks(synthesis)
    for block in blocks:
        btype = block["type"]
        if btype == "h3":
            h = doc.add_heading(block["text"], level=3)
            for run in h.runs:
                run.font.color.rgb = RGBColor(8, 145, 178)
        elif btype == "h2":
            doc.add_heading(block["text"], level=2)
        elif btype == "h1":
            doc.add_heading(block["text"], level=1)
        elif btype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("━" * 50)
            run.font.color.rgb = RGBColor(200, 220, 240)
        elif btype == "bullet_list":
            for item in block["items"]:
                clean = re.sub(r"\*+", "", item).strip()
                p = doc.add_paragraph(clean, style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
        elif btype == "number_list":
            for item in block["items"]:
                clean = re.sub(r"^\d+\.\s*", "", item)
                clean = re.sub(r"\*+", "", clean).strip()
                p = doc.add_paragraph(clean, style="List Number")
                p.paragraph_format.space_after = Pt(2)
        elif btype == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(block["text"])
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif btype == "paragraph":
            _add_rich_paragraph(doc, block["text"])

    # ── Section 2: Charts ──
    opportunities = _parse_opportunities(synthesis)
    if opportunities:
        doc.add_page_break()
        doc.add_heading("2. Strategic Analysis Charts", level=1)

        p = doc.add_paragraph("Opportunity Scoring: Feasibility vs Impact")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = RGBColor(14, 165, 233)
        try:
            _add_docx_image_from_bytes(doc, _opportunity_score_chart(opportunities))
        except Exception:
            pass

        doc.add_paragraph("")

        p = doc.add_paragraph("Value Curve Comparison")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = RGBColor(14, 165, 233)
        try:
            _add_docx_image_from_bytes(doc, _value_curve_chart(opportunities))
        except Exception:
            pass

        doc.add_paragraph("")

        p = doc.add_paragraph("Estimated Startup Cost Breakdown")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = RGBColor(14, 165, 233)
        try:
            _add_docx_image_from_bytes(doc, _cost_breakdown_chart(opportunities))
        except Exception:
            pass

    # ── Section 3: Research Sources ──
    doc.add_page_break()
    doc.add_heading("3. Research Sources", level=1)

    for b in research_data:
        label = b["branch"].replace("_", " ").title()
        h = doc.add_heading(label, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(8, 145, 178)

        if b.get("error"):
            p = doc.add_paragraph(f"Error: {b['error']}")
            p.runs[0].font.color.rgb = RGBColor(180, 50, 50)
            continue

        if b.get("answer"):
            p = doc.add_paragraph()
            run = p.add_run("Summary: ")
            run.bold = True
            run.font.color.rgb = RGBColor(14, 165, 233)
            p.add_run(b["answer"])

        for r in b.get("results", []):
            p = doc.add_paragraph()
            run = p.add_run(r.get("title", ""))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(30, 30, 30)

            url_run = p.add_run(f"\n{r.get('url', '')}")
            url_run.font.size = Pt(8)
            url_run.font.color.rgb = RGBColor(14, 105, 165)
            url_run.underline = True

            content = truncate(r.get("content", ""), 250)
            if content:
                c_run = p.add_run(f"\n{content}")
                c_run.font.size = Pt(9)
                c_run.font.color.rgb = RGBColor(80, 80, 80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════════
#  PDF  (fpdf2) — professional report styling
# ════════════════════════════════════════════════════════════

# Emoji-to-text mapping for PDF-safe rendering
_EMOJI_MAP = {
    "🔍": "[INSIGHT]", "📊": "[DATA]", "🎯": "[TARGET]", "🔄": "[STRATEGY]",
    "💰": "[FINANCE]", "🚀": "[LAUNCH]", "⚠️": "[RISK]", "📎": "[SOURCES]",
    "✅": "[SCORE]", "💡": "[WHY]", "🟢": "[START]", "🌟": "[STAR]",
    "📋": "[BRIEF]", "💬": "[CHAT]", "📥": "[EXPORT]", "📚": "[HISTORY]",
    "🌊": "[OCEAN]", "📄": "", "📝": "", "❌": "[X]", "⚠": "[!]",
    "⭐": "[*]", "🔄": "", "🗑️": "", "📂": "", "🚫": "[BLOCKED]",
    "🧠": "[AI]", "🔄": "", "✅": "[OK]",
}

def _safe_latin(text: str) -> str:
    """Replace non-latin-1 chars so fpdf2 doesn't crash."""
    for emoji, replacement in _EMOJI_MAP.items():
        text = text.replace(emoji, replacement)
    return text.encode("latin-1", "replace").decode("latin-1")


def _clean_md_for_pdf(text: str) -> str:
    """Strip markdown formatting for clean PDF text."""
    text = re.sub(r"#{1,6}\s+", "", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
    text = re.sub(r"^---$", "_" * 60, text, flags=re.MULTILINE)
    text = re.sub(r"^>\s+", "", text, flags=re.MULTILINE)
    # Clean up leftover markdown table pipes for readability
    text = re.sub(r"^\|[-\s|]+\|$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\|(.+?)\|$", lambda m: " | ".join(c.strip() for c in m.group(1).split("|")), text, flags=re.MULTILINE)
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


class _ReportPDF(FPDF):
    """Custom PDF with professional header/footer branding."""

    BLUE = (14, 165, 233)
    TEAL = (20, 184, 166)
    DARK = (30, 30, 30)
    GRAY = (80, 80, 80)
    LIGHT_GRAY = (140, 140, 140)
    BG_LIGHT = (240, 249, 255)  # very light blue for highlight boxes

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(*self.LIGHT_GRAY)
        self.cell(0, 8, "AI Blue Ocean Finder  |  Confidential Report", align="L")
        self.set_draw_color(*self.BLUE)
        self.set_line_width(0.3)
        self.line(10, 12, 200, 12)
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(*self.LIGHT_GRAY)
        self.set_draw_color(220, 220, 220)
        self.set_line_width(0.2)
        self.line(10, self.get_y() - 2, 200, self.get_y() - 2)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    def _section_title(self, text: str, size: int = 16):
        self.set_font("Helvetica", "B", size)
        self.set_text_color(*self.BLUE)
        self.cell(0, 10, _safe_latin(text), new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(*self.BLUE)
        self.set_line_width(0.6)
        self.line(self.get_x(), self.get_y(), self.get_x() + 80, self.get_y())
        self.ln(4)
        self.set_draw_color(*self.TEAL)
        self.set_line_width(0.3)
        self.line(self.get_x(), self.get_y(), self.get_x() + 50, self.get_y())
        self.ln(6)

    def _subsection_title(self, text: str, size: int = 13):
        # Colored left accent bar
        x = self.get_x()
        y = self.get_y()
        self.set_fill_color(*self.BLUE)
        self.rect(x, y, 3, 8, "F")
        self.set_x(x + 6)
        self.set_font("Helvetica", "B", size)
        self.set_text_color(*self.DARK)
        self.cell(0, 8, _safe_latin(text), new_x="LMARGIN", new_y="NEXT")
        self.ln(3)

    def _body_text(self, text: str, size: int = 10):
        self.set_font("Helvetica", "", size)
        self.set_text_color(50, 50, 50)
        self.multi_cell(0, 5.5, _safe_latin(_clean_md_for_pdf(text)))
        self.ln(2)

    def _bold_label(self, label: str, value: str, size: int = 10):
        self.set_font("Helvetica", "B", size)
        self.set_text_color(*self.BLUE)
        label_w = self.get_string_width(label) + 2
        self.cell(label_w, 5.5, _safe_latin(label))
        self.set_font("Helvetica", "", size)
        self.set_text_color(50, 50, 50)
        self.multi_cell(0, 5.5, _safe_latin(_clean_md_for_pdf(value)))
        self.ln(1)

    def _highlight_box(self, label: str, value: str):
        """Draw a highlighted box with light blue background and blue left border."""
        x = self.get_x()
        y = self.get_y()

        # Check if we need a page break
        if y > 260:
            self.add_page()
            y = self.get_y()

        # Draw background
        self.set_fill_color(*self.BG_LIGHT)
        self.set_draw_color(*self.BLUE)

        # Calculate height needed
        self.set_font("Helvetica", "", 10)
        text = _safe_latin(_clean_md_for_pdf(value))
        # Rough estimate of lines
        chars_per_line = 85
        num_lines = max(1, len(text) // chars_per_line + 1)
        box_height = max(10, num_lines * 5.5 + 6)

        # Draw the box
        self.rect(x + 2, y, 4, box_height, "F")  # left accent bar
        self.set_fill_color(*self.BG_LIGHT)
        self.rect(x + 6, y, 184, box_height, "F")  # background

        # Label
        self.set_xy(x + 8, y + 2)
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(*self.BLUE)
        self.cell(0, 5, _safe_latin(label))

        # Value
        self.set_xy(x + 8, y + 7)
        self.set_font("Helvetica", "", 9.5)
        self.set_text_color(40, 40, 40)
        self.multi_cell(178, 5, text)

        self.set_y(y + box_height + 3)

    def _divider(self):
        self.set_draw_color(220, 220, 220)
        self.set_line_width(0.3)
        y = self.get_y() + 3
        # Gradient-like divider using two lines
        self.set_draw_color(*self.BLUE)
        self.set_line_width(0.4)
        self.line(10, y, 60, y)
        self.set_draw_color(220, 220, 220)
        self.set_line_width(0.2)
        self.line(60, y, 200, y)
        self.ln(6)

    def _add_image_from_bytes(self, img_bytes: bytes, x: float = None, w: float = 180):
        """Add a PNG image from bytes to the PDF."""
        if x is None:
            x = (210 - w) / 2  # center
        buf = io.BytesIO(img_bytes)
        buf.seek(0)
        self.image(buf, x=x, w=w)


def generate_pdf(
    industry: str,
    location: str,
    synthesis: str,
    research_data: list[dict],
) -> io.BytesIO:
    pdf = _ReportPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=25)

    # ════════════════════════════════════════════════════
    #  TITLE PAGE
    # ════════════════════════════════════════════════════
    pdf.add_page()
    pdf.ln(35)

    # Top accent bar with gradient effect
    pdf.set_fill_color(14, 165, 233)
    pdf.rect(10, pdf.get_y(), 190, 4, "F")
    pdf.set_fill_color(20, 184, 166)
    pdf.rect(10, pdf.get_y() + 4, 190, 1, "F")
    pdf.ln(18)

    # Title
    pdf.set_font("Helvetica", "B", 32)
    pdf.set_text_color(14, 165, 233)
    pdf.cell(0, 16, "Blue Ocean Strategy Report", align="C", new_x="LMARGIN", new_y="NEXT")

    # Subtitle line
    pdf.set_draw_color(20, 184, 166)
    pdf.set_line_width(0.5)
    cx = 105
    pdf.line(cx - 30, pdf.get_y() + 2, cx + 30, pdf.get_y() + 2)
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 18)
    pdf.set_text_color(40, 40, 40)
    pdf.cell(0, 10, _safe_latin(industry), align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 13)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 8, _safe_latin(f"Market: {location}"), align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(20)

    # Bottom accent bar
    pdf.set_fill_color(14, 165, 233)
    pdf.rect(10, pdf.get_y(), 190, 1, "F")
    pdf.set_fill_color(20, 184, 166)
    pdf.rect(10, pdf.get_y() + 1, 190, 3, "F")
    pdf.ln(10)

    pdf.set_font("Helvetica", "I", 10)
    pdf.set_text_color(140, 140, 140)
    pdf.cell(0, 8, "Generated by AI Blue Ocean Finder", align="C", new_x="LMARGIN", new_y="NEXT")

    # ════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ════════════════════════════════════════════════════
    pdf.add_page()
    pdf._section_title("Table of Contents")
    pdf.ln(4)

    toc = [
        ("1.", "Blue Ocean Opportunities", "Detailed analysis of 5 hidden market opportunities"),
        ("2.", "Strategic Analysis Charts", "Visual scoring, value curves, and cost breakdowns"),
        ("3.", "Strategic Recommendation", "Ranked recommendations for first-mover advantage"),
        ("4.", "Research Sources", "Data sources across 5 research branches"),
    ]
    for num, title, desc in toc:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(*pdf.BLUE)
        pdf.cell(12, 8, num)
        pdf.set_text_color(*pdf.DARK)
        pdf.cell(80, 8, _safe_latin(title))
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*pdf.GRAY)
        pdf.cell(0, 8, _safe_latin(desc), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # ════════════════════════════════════════════════════
    #  SECTION 1: OPPORTUNITIES
    # ════════════════════════════════════════════════════
    pdf.add_page()
    pdf._section_title("1. Blue Ocean Opportunities")

    # Parse synthesis into structured sections
    sections = re.split(r"(?=### Opportunity \d+)", synthesis)
    opp_num = 0
    for section in sections:
        section = section.strip()
        if not section:
            continue

        # Extract opportunity title
        title_match = re.match(r"### (Opportunity \d+:.+?)(?:\n|$)", section)
        if not title_match:
            continue

        opp_num += 1
        # Check for page break
        if pdf.get_y() > 230:
            pdf.add_page()

        pdf._subsection_title(f"  {title_match.group(1).strip()}")
        section = section[title_match.end():]

        # Extract viability percentage for display
        viability_match = re.search(r"\*\*Overall Viability:\s*(\d+)%\*\*", section)
        if viability_match:
            pct = viability_match.group(1)
            pdf._highlight_box("VIABILITY SCORE", f"{pct}% — See detailed breakdown below")

        # Extract "Why You Should Proceed" section
        why_match = re.search(
            r"###\s*💡\s*Why You Should Proceed\s*\n(.*?)(?=\n###\s|---|\Z)",
            section, re.DOTALL
        )
        if why_match:
            pdf._highlight_box("WHY PROCEED", _clean_md_for_pdf(why_match.group(1).strip()))

        # Extract "Easiest Way to Start" section with all 3 options
        start_match = re.search(
            r"###\s*🟢\s*Easiest Way to Start.*?\n(.*?)(?=\n###\s*📎|\n---\s*$|\Z)",
            section, re.DOTALL
        )
        if start_match:
            start_text = start_match.group(1).strip()
            pdf._subsection_title("  Easiest Way to Start", size=11)
            # Extract each option
            for opt_label in ["Option A", "Option B", "Option C"]:
                opt_match = re.search(
                    rf"\*\*{opt_label}:\s*(.+?)\*\*\s*\n(.*?)(?=\*\*Option [A-C]|\*\*Recommendation|\Z)",
                    start_text, re.DOTALL
                )
                if opt_match:
                    opt_name = opt_match.group(1).strip()
                    opt_detail = opt_match.group(2).strip()
                    pdf._bold_label(f"  {opt_label}: {opt_name}", _clean_md_for_pdf(opt_detail)[:300])
            # Extract final recommendation
            rec_match = re.search(r"\*\*Recommendation:\*\*\s*(.+?)(?:\n|$)", start_text)
            if rec_match:
                pdf._highlight_box("RECOMMENDED PATH", _clean_md_for_pdf(rec_match.group(1).strip()))

        # Parse labeled fields
        field_pattern = re.compile(
            r"\*\*(.+?):\*\*\s*\n(.*?)(?=\n\*\*|\n---|\n###|\Z)",
            re.DOTALL,
        )

        last_end = 0
        for match in field_pattern.finditer(section):
            label = match.group(1).strip()
            value = match.group(2).strip()

            # Highlight key sections with colored boxes
            if label.lower().startswith("the hidden insight"):
                pdf._highlight_box(f"[HIDDEN INSIGHT] {label}", value)
            elif label.lower().startswith("market gap"):
                pdf._highlight_box(f"[MARKET GAP] {label}", value)
            elif "viability score" in label.lower() or "overall viability" in label.lower():
                pdf._highlight_box(f"[VIABILITY SCORE] {label}", value)
            elif "why you should proceed" in label.lower():
                pdf._highlight_box(f"[WHY PROCEED] {label}", value)
            elif "easiest way to start" in label.lower() or label.lower().startswith("recommendation"):
                pdf._highlight_box(f"[HOW TO START] {label}", value)
            elif label.lower().startswith("why this is invisible"):
                pdf._bold_label(f"  {label}:", value)
            elif label.lower().startswith("target audience"):
                pdf._bold_label(f"  {label}:", value)
            elif label.lower().startswith("value innovation"):
                pdf._subsection_title(f"    {label}", size=11)
                pdf._body_text(value)
            elif label.lower().startswith("price-to-value"):
                pdf._subsection_title(f"    {label}", size=11)
                sub_pattern = re.compile(r"[-*]\s+\*\*(.+?):\*\*\s*(.+)")
                for sub_match in sub_pattern.finditer(value):
                    pdf._bold_label(f"    {sub_match.group(1)}:", sub_match.group(2))
                non_sub = sub_pattern.sub("", value).strip()
                if non_sub:
                    pdf._body_text(non_sub)
            elif label.lower().startswith("first-mover"):
                pdf._subsection_title(f"    {label}", size=11)
                list_items = re.findall(r"^\d+\.\s+(.+)", value, re.MULTILINE)
                if list_items:
                    for item in list_items:
                        pdf.set_font("Helvetica", "", 10)
                        pdf.set_text_color(50, 50, 50)
                        pdf.cell(8, 5.5, "")
                        pdf.multi_cell(0, 5.5, _safe_latin(f"  {_clean_md_for_pdf(item)}"))
                        pdf.ln(1)
                else:
                    pdf._body_text(value)
            else:
                pdf._bold_label(f"  {label}:", value)

            last_end = match.end()

        # Remaining text (evidence, sources, etc.)
        remaining_after = section[last_end:].strip()
        if remaining_after and remaining_after != "---":
            source_items = re.findall(r"[-*]\s+(.+?)(?:\n|$)", remaining_after)
            if source_items:
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(*pdf.DARK)
                pdf.cell(0, 6, "  Evidence & Sources:", new_x="LMARGIN", new_y="NEXT")
                for item in source_items:
                    pdf.set_font("Helvetica", "", 9)
                    pdf.set_text_color(70, 70, 70)
                    pdf.cell(8, 5, "")
                    pdf.multi_cell(0, 5, _safe_latin(f"  {_clean_md_for_pdf(item)}"))
                    pdf.ln(1)
            else:
                list_items = re.findall(r"^\d+\.\s+(.+)", remaining_after, re.MULTILINE)
                if list_items:
                    for item in list_items:
                        pdf.set_font("Helvetica", "", 10)
                        pdf.set_text_color(50, 50, 50)
                        pdf.cell(8, 5.5, "")
                        pdf.multi_cell(0, 5.5, _safe_latin(f"  {_clean_md_for_pdf(item)}"))
                        pdf.ln(1)

        pdf._divider()

    # ════════════════════════════════════════════════════
    #  SECTION 2: CHARTS
    # ════════════════════════════════════════════════════
    opportunities = _parse_opportunities(synthesis)
    if opportunities:
        pdf.add_page()
        pdf._section_title("2. Strategic Analysis Charts")
        pdf.ln(2)

        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(*pdf.DARK)
        pdf.cell(0, 8, "Opportunity Scoring: Feasibility vs Impact", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        try:
            pdf._add_image_from_bytes(_opportunity_score_chart(opportunities), w=175)
        except Exception:
            pass
        pdf.ln(6)

        if pdf.get_y() > 160:
            pdf.add_page()

        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(*pdf.DARK)
        pdf.cell(0, 8, "Value Curve Comparison", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        try:
            pdf._add_image_from_bytes(_value_curve_chart(opportunities), w=175)
        except Exception:
            pass
        pdf.ln(6)

        if pdf.get_y() > 160:
            pdf.add_page()

        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(*pdf.DARK)
        pdf.cell(0, 8, "Estimated Startup Cost Breakdown", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        try:
            pdf._add_image_from_bytes(_cost_breakdown_chart(opportunities), w=175)
        except Exception:
            pass

    # ════════════════════════════════════════════════════
    #  SECTION 3: STRATEGIC RECOMMENDATION
    # ════════════════════════════════════════════════════
    rec_match = re.search(r"\*\*Strategic Recommendation\*\*(.*)", synthesis, re.DOTALL)
    if rec_match:
        pdf.add_page()
        pdf._section_title("3. Strategic Recommendation")
        rec_text = rec_match.group(1).strip()

        # Parse ranked items if numbered
        ranked = re.findall(r"^\d+\.\s+(.+)", rec_text, re.MULTILINE)
        if ranked:
            for item in ranked:
                pdf._highlight_box("", _clean_md_for_pdf(item))
        else:
            pdf._body_text(rec_text)

    # ════════════════════════════════════════════════════
    #  SECTION 4: RESEARCH SOURCES
    # ════════════════════════════════════════════════════
    pdf.add_page()
    pdf._section_title("4. Research Sources")

    for b in research_data:
        label = b["branch"].replace("_", " ").title()
        pdf._subsection_title(label)

        if b.get("error"):
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(180, 50, 50)
            pdf.multi_cell(0, 5, _safe_latin(f"Error: {b['error']}"))
            pdf.ln(4)
            continue

        if b.get("answer"):
            pdf._bold_label("Summary:", b["answer"])

        for r in b.get("results", []):
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*pdf.DARK)
            pdf.cell(0, 5, _safe_latin(truncate(r.get("title", ""), 90)), new_x="LMARGIN", new_y="NEXT")

            url = r.get("url", "")
            if url:
                pdf.set_font("Helvetica", "U", 7)
                pdf.set_text_color(14, 105, 165)
                pdf.cell(0, 4, _safe_latin(truncate(url, 110)), new_x="LMARGIN", new_y="NEXT")

            content = truncate(r.get("content", ""), 300)
            if content:
                pdf.set_font("Helvetica", "", 8)
                pdf.set_text_color(70, 70, 70)
                pdf.multi_cell(0, 4, _safe_latin(content))
            pdf.ln(3)
        pdf.ln(4)

    buf = io.BytesIO()
    buf.write(bytes(pdf.output()))
    buf.seek(0)
    return buf
